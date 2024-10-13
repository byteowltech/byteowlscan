import base64
from io import BytesIO

import fitz  # PyMuPDF
import mammoth
from docx import Document
from markdownify import markdownify as md
from pdfminer.high_level import extract_text_to_fp
from semantic_text_splitter import MarkdownSplitter, TextSplitter
from tokenizers import Tokenizer


def initialize_text_splitter(max_tokens=16384):
    """
    Initializes a text splitter for generic text content.

    Args:
        max_tokens (int): Maximum number of tokens in each chunk.

    Returns:
        TextSplitter: Initialized text splitter.
    """
    tokenizer = Tokenizer.from_pretrained("bert-base-uncased")
    return TextSplitter.from_huggingface_tokenizer(tokenizer, max_tokens)

def initialize_markdown_splitter(max_tokens=16384):
    """
    Initializes a text splitter specifically for Markdown content.

    Args:
        max_tokens (int): Maximum number of tokens in each chunk.

    Returns:
        MarkdownSplitter: Initialized Markdown text splitter.
    """
    tokenizer = Tokenizer.from_pretrained("bert-base-uncased")
    return MarkdownSplitter.from_huggingface_tokenizer(tokenizer, max_tokens)

def convert_image(image):
    """
    Placeholder function to convert an image to base64 or other formats for HTML.

    Args:
        image: The image to be converted.

    Returns:
        str: Empty string as a placeholder.
    """
    return ""

def convert_docx_to_html(docx_path):
    """
    Converts a DOCX file to HTML content without images.

    Args:
        docx_path (str): Path to the DOCX file.

    Returns:
        str: HTML content without images.
    """
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(convert_image))
        html_content = result.value.replace("<img />", "")
    return html_content

def convert_html_to_markdown(html_content):
    """
    Converts HTML content to Markdown format.

    Args:
        html_content (str): HTML content.

    Returns:
        str: Markdown content.
    """
    return md(html_content)

def convert_pdf_to_html(file_path):
    """
    Converts a PDF file to HTML content.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        str: HTML content extracted from the PDF.
    """
    output_buffer = BytesIO()
    with open(file_path, 'rb') as pdf_file:
        extract_text_to_fp(pdf_file, output_buffer, output_type='html')
    return output_buffer.getvalue().decode('utf-8')

def read_docx_to_text(filename):
    """
    Reads text from a DOCX file, including text from paragraphs and tables.

    Args:
        filename (str): Path to the DOCX file.

    Returns:
        str: Extracted text from the DOCX file.
    """
    doc = Document(filename)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text:
            full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            full_text.append('\t'.join(row_data))
    
    return '\n'.join(full_text)

def split_markdown_into_chunks(content, max_tokens=16384):
    """
    Splits Markdown content into smaller chunks based on the token limit.

    Args:
        content (str): Markdown content to be split.
        max_tokens (int): Maximum number of tokens in each chunk.

    Returns:
        list: List of content chunks.
    """
    splitter = initialize_markdown_splitter(max_tokens)
    return splitter.chunks(content)

def split_text_into_chunks(content, max_tokens=16384):
    """
    Splits generic text content into smaller chunks based on the token limit.

    Args:
        content (str): Text content to be split.
        max_tokens (int): Maximum number of tokens in each chunk.

    Returns:
        list: List of content chunks.
    """
    splitter = initialize_text_splitter(max_tokens)
    return splitter.chunks(content)
